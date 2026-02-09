import streamlit as st
import openai
from docx import Document
from docx.shared import RGBColor
from io import BytesIO
import re

# --- 1. CONFIGURATION ---
st.set_page_config(page_title="Productivity Strategy Auto-Filler", layout="wide")

# Check for API Key
if "OPENAI_API_KEY" in st.secrets:
    openai_api_key = st.secrets["OPENAI_API_KEY"]
else:
    openai_api_key = None

# --- 2. SECTIONS MAPPING ---
SECTIONS_MAPPING = [
# --- SECTION 1 ---
    {
        "anchor_text": "Focus: Business Section", 
        "prompt": """
        **Task:** Write a formal executive summary for the "Focus: Business Section".
        
        **STYLE GUIDELINES (Strict):**
        1. **Format:** Use continuous narrative paragraphs. **DO NOT** use Markdown headers (like ## or ###).
        2. **Tone:** Professional, objective, and concise. Do not use phrases like "The transcript says" or "Jamie mentions." State the facts directly.
        3. **Structure:**
        - **Para 1 (Context):** Summarize the company's operations, key clients (e.g., Vivobarefoot, Army), and background. (Include address/registration details only if present in text).
        - **Para 2 (Strategy):** Explain the commercial priority (e.g., reliance on Vivobarefoot, goal to improve efficiency/margins).
        - **Para 3 (The Focus):** Define the chosen area for productivity improvement (e.g., the resole process) and why it is critical.
        - **Ending List:** Conclude with a simple text list of the top-level productivity focus areas (e.g., Direct labour productivity, Throughput rate).
        """
    },
    
    # --- SECTION 2 ---
    {
        "anchor_text": "Measure: Business Section", 
        "prompt": """
        **Main Context:** Identification and alignment of relevant KPI/productivity metrics.
        
        **Please answer these specific sub-questions:**
        1. What is the current level of output/throughput/capacity?
        2. What is the high-level productivity measure (e.g., Total productivity, Labour productivity)?
        3. What is the current level of this chosen productivity metric?
        4. Identify relevant business process KPIs that align/affect your chosen productivity focus area.
        5. If relevant data is not available, how will you plan to collect it in the future?
        6. What is the management process/plan for ongoing extraction and monitoring of key data sets?
        """
    },
    
    # --- SECTION 3 ---
    {
        "anchor_text": "Focus: Improvement", # The code will normalize this to 'focus:improvement'
        "prompt": """
        **Main Context:** An outline of the improvement plan for the focus area.
        
        **Please answer these specific sub-questions:**
        1. Identify the main [root] causes for the current KPI performance.
        2. Describe the improvement plan to address these causes and justify how it will improve high-level productivity.
        3. What is the expected performance improvement in your business process KPI?
        4. What is the expected new level of output/throughput/capacity?
        5. What is the expected change in the high-level productivity measure?
        6. What is the expected benefit and/or financial savings?
        """
    },


# --- SECTION 4: ACTIONS & MILESTONES (NEW) ---
    {
        "anchor_text": "Focus:Action/Milestone", 
        "prompt": """
        **Task:** Extract 3 high-priority process improvement actions.
        
        **STYLE GUIDELINES (Strict):**
        1. **Tone:** "Process Engineer" / "Lean Six Sigma" style. 
           - Use imperative verbs to start every line (e.g., "Complete", "Define", "Map", "Implement").
           - Use technical terms if relevant (e.g., SOPs, Value Stream Map, Skills Matrix, FTEs).
        2. **Format:** A clean, vertical list of punchy phrases. 
           - **DO NOT** use "Action:" or "Milestone:" labels. 
           - **DO NOT** write long explanations. Keep it under 15 words per item.
        3. **Content:** Focus on the practical steps to execute the plan (e.g., mapping the process, defining roles, collecting data).
        """
    },

# --- SECTION 5: ACTIONS & MILESTONES (NEW) ---
    {
        "anchor_text": "Measure:Action/Milestone", 
        "prompt": """
        **Task:** List 3 key measurement frameworks and resource controls to be implemented.
        
        **STYLE GUIDELINES:**
        1. **Format:** A concise vertical list.
        2. **Content Requirements:**
           - Mention "Selection and implementation of KPIs by category" (Outcome, Quality, Flow, Cost, Delivery).
           - Mention FTE utilisation or workforce tracking.
           - Mention Physical changes (e.g., Layout, Automation, Storage).
           - Mention Incentives (e.g., Bonus schemes, RFT targets).
        3. **Tone:** Operational and structural.
        4. **Format:** A clean, vertical list of punchy phrases. 
    -       -DO NOT** use "Action:" or "Milestone:" labels. 
            -DO NOT** write long explanations. Keep it under 15 words per item.
            - Keep to 3 points.
        """
    },

# --- SECTION 6: ACTIONS & MILESTONES (NEW) ---
    {
        "anchor_text": "Improve:Action/Milestone", 
        "prompt": """
        **Task:** List 3 specific data metrics and productivity formulas to be monitored.
        
        **STYLE GUIDELINES:**
        1. **Format:** A concise vertical list.
        2. **Content Requirements:**
           - Mention "Time by process" (variations, limits).
           - Mention "Direct labour productivity" (Sales/Costs breakdown).
           - Mention "Throughput" and "Process Flow".
           - Mention "Rework" and "Root Cause Identification".
        3. **Tone:** Analytical and data-focused.
        4. **Format:** A clean, vertical list of punchy phrases. 
    -       - DO NOT** use "Action:" or "Milestone:" labels. 
            - DO NOT** write long explanations. Keep it under 15 words per item.
            - Keep to 3 points.
        """
    }
]

# --- 3. HELPER FUNCTIONS ---

def normalize_text(text):
    """
    Removes all whitespace and converts to lowercase.
    Example: "Focus:   Business Section" -> "focus:businesssection"
    """
    if not text:
        return ""
    # Remove all whitespace characters (spaces, tabs, newlines)
    return re.sub(r'\s+', '', text).lower()

def extract_text(uploaded_file):
    """Reads text from file for the AI context."""
    if uploaded_file.type == "text/plain":
        return str(uploaded_file.read(), "utf-8")
    elif "wordprocessingml" in uploaded_file.type:
        doc = Document(uploaded_file)
        return '\n'.join([p.text for p in doc.paragraphs])
    return None

def generate_ai_response(client, transcript, prompt):
    """Generates the AI answer."""
    full_prompt = f"""
    You are a Senior Business Analyst.
    TRANSCRIPT CONTEXT:
    {transcript[:30000]}
    
    QUESTION TO ANSWER:
    {prompt}
    
    INSTRUCTIONS:
    - Provide a comprehensive, detailed narrative answer.
    - Base your answer strictly on the transcript provided.
    """
    try:
        # Use gpt-5.2 (or gpt-4o if 5.2 unavailable)
        response = client.chat.completions.create(
            model="gpt-5.2", 
            messages=[
                {
                    "role": "system", 
                    "content": (
                        "You are a Senior Business Analyst writing a formal strategic report. "
                        "Your output must be a cohesive narrative, not a Q&A list. "
                        "Use professional, third-person business language (e.g., 'The company operates...', 'The strategic focus is...'). "
                        "Avoid conversational fillers and do not use Markdown headers."
                    )
                },
                {"role": "user", "content": full_prompt}
            ],
            temperature=0.3 # Lower temperature for more consistent, factual writing
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Error: {e}"

def fill_word_template(template_file, transcript_text, api_key):
    """
    Locates headers in Tables using 'fuzzy' matching and fills the cell below.
    """
    client = openai.OpenAI(api_key=api_key)
    doc = Document(template_file)
    
    progress_bar = st.progress(0)
    status_text = st.empty()
    total_steps = len(SECTIONS_MAPPING)

    # Debug logs to help user see what was found
    debug_logs = []

    for i, section in enumerate(SECTIONS_MAPPING):
        raw_anchor = section["anchor_text"]
        target_anchor = normalize_text(raw_anchor) # e.g. "focus:businesssection"
        prompt = section["prompt"]
        
        status_text.text(f"Looking for header: '{raw_anchor}'...")
        
        # 1. Generate Answer
        ai_answer = generate_ai_response(client, transcript_text, prompt)
        
        found = False
        
        # 2. Search Tables
        for table_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    
                    # Normalize cell text for comparison
                    cell_clean = normalize_text(cell.text)
                    
                    # Check if our target text is inside this cell
                    if target_anchor in cell_clean:
                        debug_logs.append(f"✅ Found '{raw_anchor}' in Table {table_idx+1}, Row {r_idx+1}")
                        
                        # We found the header! Now try to write to the cell BELOW it.
                        try:
                            # Strategy: Look at the next row (r_idx + 1)
                            if r_idx + 1 < len(table.rows):
                                target_cell = table.cell(r_idx + 1, c_idx)
                                
                                # Clear placeholder text if any
                                target_cell.text = "" 
                                
                                # Add the answer
                                p = target_cell.paragraphs[0]
                                run = p.add_run(ai_answer)
                                run.font.color.rgb = RGBColor(0, 50, 100)
                                run.font.name = 'Calibri'
                                found = True
                            else:
                                debug_logs.append(f"⚠️ Found header but no row exists below it in Table {table_idx+1}")
                        except Exception as e:
                            debug_logs.append(f"❌ Error writing to cell: {e}")
                        
                        if found: break
                if found: break
            if found: break

        if not found:
            debug_logs.append(f"❌ Could not find header '{raw_anchor}' anywhere in the tables.")

        progress_bar.progress((i + 1) / total_steps)

    status_text.text("Processing Complete!")
    
    # Save result
    output_stream = BytesIO()
    doc.save(output_stream)
    output_stream.seek(0)
    return output_stream, debug_logs

# --- 4. UI ---
st.title("🏭 Productivity Strategy Auto-Filler (Fix)")
st.markdown("This version ignores extra spaces/tabs to find your headers reliably.")

# Sidebar
with st.sidebar:
    st.header("Settings")
    if openai_api_key:
        st.success("API Key Loaded", icon="✅")
    else:
        openai_api_key = st.text_input("API Key:", type="password")

# Uploads
col1, col2 = st.columns(2)
with col1:
    transcript_file = st.file_uploader("1. Upload Transcript", type=["txt", "docx"])
with col2:
    template_file = st.file_uploader("2. Upload Template", type=["docx"])

# Run
if st.button("Generate Report", type="primary"):
    if not openai_api_key or not transcript_file or not template_file:
        st.error("Missing files or API key.")
    else:
        with st.spinner("Reading transcript..."):
            transcript_text = extract_text(transcript_file)
        
        if transcript_text:
            with st.spinner("Filling Template..."):
                try:
                    final_doc, logs = fill_word_template(template_file, transcript_text, openai_api_key)
                    
                    st.success("Report Generated!")
                    
                    # Download
                    st.download_button(
                        label="📥 Download Result",
                        data=final_doc,
                        file_name="Completed_Strategy_Plan.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    
                    # Debug Info
                    with st.expander("Show Processing Logs (Debug)"):
                        for log in logs:
                            st.write(log)
                            
                except Exception as e:
                    st.error(f"Critical Error: {e}")
